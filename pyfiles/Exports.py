# -*- coding: utf-8 -*-

import pandas as pd
import numpy as pd
from PIL import Image
from docx.shared import Inches
from docx.shared import Pt
from docx import Document
from docx.shared import RGBColor
from Main import trimestre as Trimestre, año as Año
from Main import MM_Ing_STR, mesesSTR, MM_Demanda_STR, MM_Cargos, MesesADD, AñosADD
from Main import ConvList, ingSTNString, cargosTString, contString
from tkinter import *
import os


## Valores predeterminados de inputs Analistas
AnalistaCargosSTN1 = "XXXXXXX"
AnalistaCargosSTN2 = "XXXXXXX"
AnalistaCargosSTN3 = "XXXXXXX"
AnalistaIngSTN1 = "XXXXXXX"
AnalistaIngSTN2 = "XXXXXXX"
AnalistaIngSTN3 = "XXXXXXX"
AnalistaContribSTN1 = "XXXXXXX"
AnalistaContribSTN2 = "XXXXXXX"
AnalistaContribSTN3 = "XXXXXXX"


## Interfaz de datos
def CrearStringsAnalistas():
    global AnalistaCargosSTN1
    global AnalistaCargosSTN2
    global AnalistaCargosSTN3
    global AnalistaIngSTN1 
    global AnalistaIngSTN2 
    global AnalistaIngSTN3 
    global AnalistaContribSTN1
    global AnalistaContribSTN2 
    global AnalistaContribSTN3 
    AnalistaCargosSTN1 = CargosSTN1_input.get()
    AnalistaCargosSTN2 = CargosSTN2_input.get()
    AnalistaCargosSTN3 = CargosSTN3_input.get()
    AnalistaIngSTN1 = IngSTN1_input.get()
    AnalistaIngSTN2 = IngSTN2_input.get()
    AnalistaIngSTN3 = IngSTN3_input.get()
    AnalistaContribSTN1 = ContribSTN1_input.get()
    AnalistaContribSTN2 = ContribSTN2_input.get()
    AnalistaContribSTN3 = ContribSTN3_input.get()    
    Mainframe.destroy()
    

Mainframe=Tk() 
Mainframe.title("Interfaz Informe Mandatario Trimestral LAC") #Titulo de la ventana
Mainframe.iconbitmap("XM-LOGO-RGB-01.ico") #Poner icono
Mainframe.geometry("600x600") #Definir el tamaño de la ventana
Mainframe.resizable(0,0) #Bloquear para que el usuario no la pueda editar

main_title = Label(text="Datos Analistas para Documento Generado",font=("Arial",13)) #Poner titulo dentro de la ventana
main_title.place(x=50,y=10) #Ubicación del texto de las variables de ingreso


# Input Analistas Cargos STN
CargosSTN_title = Label(text="Cargos en el STN",font=("Arial",13)) #Poner titulo dentro de la ventana
CargosSTN_title.place(x=70,y=40) #Ubicación del texto de las variables de ingreso

CargosSTN1_label=Label(text= "En " + mesesSTR[0] + cargosTString[0], font=("Arial",10)) 
CargosSTN1_label.place(x=70, y= 70)
CargosSTN1_input = Entry(Mainframe)
CargosSTN1_input.place(x=80, y= 90)

CargosSTN2_label=Label(text= "En " + mesesSTR[1] + cargosTString[1], font=("Arial",10)) 
CargosSTN2_label.place(x=70, y= 110)
CargosSTN2_input = Entry(Mainframe)
CargosSTN2_input.place(x=80, y= 130)

CargosSTN3_label=Label(text= "En " + mesesSTR[2] + cargosTString[2], font=("Arial",10)) 
CargosSTN3_label.place(x=70, y= 150)
CargosSTN3_input = Entry(Mainframe)
CargosSTN3_input.place(x=80, y= 170)


# Input Analistas Ingresos STN
IngSTN_title = Label(text="Ingresos en el STN",font=("Arial",13)) #Poner titulo dentro de la ventana
IngSTN_title.place(x=70,y=200) #Ubicación del texto de las variables de ingreso

IngSTN1_label=Label(text= "En " + mesesSTR[0] + ingSTNString[0], font=("Arial",10)) 
IngSTN1_label.place(x=70, y= 230)
IngSTN1_input = Entry(Mainframe)
IngSTN1_input.place(x=80, y= 250)

IngSTN2_label=Label(text= "En " + mesesSTR[1] + ingSTNString[1], font=("Arial",10)) 
IngSTN2_label.place(x=70, y= 270)
IngSTN2_input = Entry(Mainframe)
IngSTN2_input.place(x=80, y= 290)

IngSTN3_label=Label(text= "En " + mesesSTR[1] + ingSTNString[1], font=("Arial",10)) 
IngSTN3_label.place(x=70, y= 310)
IngSTN3_input = Entry(Mainframe)
IngSTN3_input.place(x=80, y= 330)
    

# Input Analistas Contribuciones STN
ContribSTN_title = Label(text="Contribuciones en el STN",font=("Arial",13)) #Poner titulo dentro de la ventana
ContribSTN_title.place(x=70,y=360) #Ubicación del texto de las variables de ingreso

ContribSTN1_label=Label(text= "En " + mesesSTR[0] + contString[0], font=("Arial",10)) 
ContribSTN1_label.place(x=70, y= 380)
ContribSTN1_input = Entry(Mainframe)
ContribSTN1_input.place(x=80, y= 400)

ContribSTN2_label=Label(text= "En " + mesesSTR[1] + contString[1], font=("Arial",10)) 
ContribSTN2_label.place(x=70, y= 420)
ContribSTN2_input = Entry(Mainframe)
ContribSTN2_input.place(x=80, y= 440)

ContribSTN3_label=Label(text= "En " + mesesSTR[2] + contString[2], font=("Arial",10)) 
ContribSTN3_label.place(x=70, y= 460)
ContribSTN3_input = Entry(Mainframe)
ContribSTN3_input.place(x=80, y= 480)


# Boton para ejecutar la funcion 
calcular_btn=Button(Mainframe,text="Guardar Datos Analista",width="30",height="2", command=CrearStringsAnalistas)
calcular_btn.place(x=300,y=550)

Mainframe.mainloop()

print("Done Interfaz!")

## Crear docmuento
try:
    os.remove("DocPruebaGenerado.docx")
except:
    print("No hay archivo, de que estas hablando")


doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Nunito'

# Funciones
def añadir_titulo(titulo, tamaño = 1):
    h = doc.add_heading(titulo, tamaño).add_run()
    font = h.font
    font.color.rgb = RGBColor(0xE6, 0x6C, 0x37)
    return h

def añadir_parrafo(parrafo, param = 0):
    if param == 0:
        line = doc.add_paragraph(parrafo)
        line.alignment = 3
    if param == 1:
        line = doc.add_paragraph(parrafo)
        line.alignment = 3
        line.style = 'List Bullet'
    if param == 2:
        line = doc.add_paragraph()
        line.alignment = 1
        runner = line.add_run(parrafo)
        runner.italic = True
    if param == 4:
        line = doc.add_paragraph()
        runner2 = line.add_run(parrafo)
        runner2.bold = True
        font = runner2.font
        font.size = Pt(12)

    if param == 5:
        line = doc.add_paragraph()
        run = doc.add_paragraph().add_run(parrafo)
        font = run.font
        font.color.rgb = RGBColor(0x44, 0x00, 0x99)
        font.size = Pt(14)
    

inputAnalistaCargosTMes1 = "No se"
inputAnalistaCargosTMes2 = "No se"
inputAnalistaCargosTMes3 = "No se"

inputAnalistaIngresosMes1 = "No se"
inputAnalistaIngresosMes2 = "No se"
inputAnalistaIngresosMes3 = "No se"

inputAnalistaContribucionesMes1 = "No se"
inputAnalistaContribucionesMes2 = "No se"
inputAnalistaContribucionesMes3 = "No se"


# 1. Parrafo Introductorio
añadir_parrafo('Liquidación y Facturación LAC', 5)

ParrafoInicial = "Las actividades principales del equipo de liquidación LAC consisten en:"
ParrafoInicialB_1 = "Cálculo de los Ingresos y Compensaciones de los Transmisores Nacionales, Liquidación y Facturación de los Cargos por Uso del Sistema de Transmisión Nacional, Liquidación y Facturación de los Servicios prestados por el LAC."
ParrafoInicialB_2 = "Liquidación y Facturación de las Contribuciones PRONE (Programa de Normalización de Redes Eléctricas), FAER (Fondo de Apoyo Financiero para la Energización de las Zonas Rurales Interconectadas) y FOES (Fondo de Energía Social)."
ParrafoInicialB_3 = "Cálculo de los Ingresos de los Operadores de Red y Transmisores Regionales, Liquidación de los Cargos por Uso de los STR y facturación de convocatorias de los Sistemas de Transmisión Regional."
ParrafoInicialB_4 = "Liquidación y Facturación de los Servicios prestados por el LAC."
ParrafoInicialB_5 = "Cálculo de los ingresos de los Operadores de Red"
ParrafoInicialB_6 = "Liquidación de Ingresos de los Operadores de Red, determinación de cargos unificados y liquidación en el esquema de Áreas de Distribución (ADD)."
ParrafoInicialB_7 = "Cálculo paralelo de los Índices de Calidad en el SDL. Lo anterior siguiendo todos los procedimientos establecidos en la reglamentación vigente."
ParrafoInicialB_8 = "Cálculo y Liquidación de los Cargos por concepto de los planes de reducción y mantenimiento de pérdidas CPROG. Adicionalmente, la evaluación y seguimiento del cumplimiento de las metas de los planes de reducción de pérdidas."

añadir_parrafo(ParrafoInicial,0)
añadir_parrafo(ParrafoInicialB_1,1)
añadir_parrafo(ParrafoInicialB_2,1)
añadir_parrafo(ParrafoInicialB_3,1)
añadir_parrafo(ParrafoInicialB_4,1)
añadir_parrafo(ParrafoInicialB_5,1)
añadir_parrafo(ParrafoInicialB_6,1)
añadir_parrafo(ParrafoInicialB_7,1)
añadir_parrafo(ParrafoInicialB_8,1)

# Principales Sucesos en LAC
ParrafoSucesosLAC = 'Principales sucesos en Liquidación y Facturación LAC para el ' + Trimestre + ' de ' + Año + "."
añadir_parrafo(ParrafoSucesosLAC,0)
SucesosList = []
with open('Principales_sucesos.txt', 'r', encoding='utf-8') as PsucesosFile:
    for line in PsucesosFile:
        parts = line.split('\n')
        añadir_parrafo(parts[0],1)
for conv in ConvList:
    añadir_parrafo(conv, 1)


##  1.Indicadores
añadir_parrafo('Indicadores económicos', 4)
doc.add_picture(r"Temp\Indicadores.png", width = Inches(6.5))
añadir_parrafo('Grafica 1, Indicadores económicos para el ' + Trimestre + ' de ' + Año + ".", 2)

## 2.STN
añadir_parrafo('Cargos y liquidación STN', 4)
añadir_parrafo('El Sistema de Transmisión Nacional (STN) Es el Sistema Interconectado de Transmisión de energía eléctrica compuesto por el conjunto de líneas, con sus correspondientes módulos de conexión, que operan a tensiones iguales o superiores a 220 kV.')

#  Grafica Cargos T
doc.add_picture(r"Temp\STN_Cargos_T.png")
añadir_parrafo('Grafica 2, Evolución de cargo T para el '+ Trimestre + ' de ' + Año + ".", 2)
añadir_parrafo("Variaciones en los cargos por uso del STN:")
añadir_parrafo('En ' + mesesSTR[0].lower() + cargosTString[0] + AnalistaCargosSTN1,1)
añadir_parrafo('En ' + mesesSTR[1].lower() + cargosTString[1] + AnalistaCargosSTN2,1)
añadir_parrafo('En ' + mesesSTR[2].lower() + cargosTString[2] + AnalistaCargosSTN3,1)

#  Grafica Contribuciones
añadir_parrafo('El valor de las contribuciones FAER, PRONE y FOES se presentan en la Gráfica 3')
doc.add_picture(r"Temp\STN_Contribuciones.png")
añadir_parrafo('Grafica 3, Evolución de Ingresos por Contribuciones para el '+ Trimestre + ' de ' + Año + ".", 2)
añadir_parrafo('En ' + mesesSTR[0].lower() + contString[0] + AnalistaContribSTN1,1)
añadir_parrafo('En ' + mesesSTR[1].lower() + contString[1] + AnalistaContribSTN2,1)
añadir_parrafo('En ' + mesesSTR[2].lower() + contString[2] + AnalistaContribSTN3,1)

#  Grafica Ingresos sin y con Contribuciones
añadir_parrafo('El Ingreso Neto a recibir por parte de los TN (sumatoria del ingreso por cargos por uso, el ingreso por las contribuciones FAER, PRONE y FOES y las compensaciones), que corresponde al total a pagar por parte de los comercializadores, siguió el comportamiento presentado en la Gráfica 4.', 0)
doc.add_picture(r"Temp\STN_Ingresos.png")
añadir_parrafo('Grafica 4, Evolución de Ingreso NETO Transmisores Nacionales.', 2)
añadir_parrafo('En ' + mesesSTR[0].lower() + ingSTNString[0] + AnalistaIngSTN1,1)
añadir_parrafo('En ' + mesesSTR[1].lower() + ingSTNString[1] + AnalistaIngSTN2,1)
añadir_parrafo('En ' + mesesSTR[2].lower() + ingSTNString[2] + AnalistaIngSTN3,1)

# Grafica Demanda STN
doc.add_picture(r"Temp\STN_Demanda.png")
añadir_parrafo('Grafica 5, Evolución de demanda energética nacional.', 2)

## 3.STR
añadir_parrafo('Cargos y liquidación STR', 4)
añadir_parrafo('El Sistema de Transmisión Regional (STR). Es el sistema de transporte de energía eléctrica compuesto por los Activos de Conexión del OR al STN y el conjunto de líneas, equipos y subestaciones, con sus equipos asociados, que operan en el Nivel de Tensión 4.')
añadir_parrafo('En cuanto al STR, los meses de' + mesesSTR[0].lower() + ', ' + mesesSTR[1].lower() + ' y ' + mesesSTR[2].lower() + ' se liquidaron según las disposiciones establecidas en las Resoluciones CREG 015 de 2018 y CREG 157 de 2011. Los cargos del nivel de tensión 4 varían mensualmente según el comportamiento de los ingresos y la demanda. De igual manera, la liquidación de los ingresos de los Operadores de Red (OR) dependen de los cambios en los cargos de cada STR y de la demanda correspondiente del mes a liquidar. ', 1)

# . Grafica STR Cargos
añadir_parrafo('En el mes de ' + MM_Cargos.lower() + ' se presentó aumento en el cargo del STR Norte y Centro Sur por aumento en la demanda para el STR norte y STR Centro Sur a pesar del aumento en los ingresos de los Transmisores Regionales y Operadores de Red.')
doc.add_picture(r"Temp\STR_Cargos.png")
añadir_parrafo('Fig 6, Evolución de cargos en el STR para el '+ Trimestre + ' de ' + Año + ".", 2)

# . Grafica STR Ingresos
añadir_parrafo('Según lo observado en la Gráfica 7, los mayores ingresos a recibir por los Operadores de Red para el STR Norte y Centro Sur se presentan en el mes de' + MM_Ing_STR.lower() + '.')
añadir_parrafo('En la siguiente gráfica se presenta la evolución de ingresos netos de los operadores de red')
doc.add_picture(r"Temp\STR_Ingresos.png")
añadir_parrafo('Fig 7, Evolución de Ingresos a recibir por los OR para el '+ Trimestre + ' de ' + Año + ".", 2)

# . Grafica STR Demanda
doc.add_picture(r"Temp\STR_Demanda.png")
añadir_parrafo('Fig 8, Evolución de demanda energetica para el STR', 2)
añadir_parrafo('Durante el trimestre, se realizó la emisión de las facturaciones para el STN. En el STR se efectuó la facturación para el mismo periodo, por concepto de remuneración de convocatorias tanto en el STR Norte como en el STR Centro-Sur y por concepto de Servicio LAC.')

## 4.ADDs
añadir_parrafo('Liquidación ADDs', 4)

añadir_parrafo('Según lo establecido en las Resoluciones CREG 058, 068, 070 de 2008 y las nuevas disposiciones establecidas en las Resolución CREG 116, 149 de 2010 y 133 de 2013, se ha efectuado la publicación de los cargos transitorios por nivel de tensión y la respectiva liquidación de las Áreas de Distribución (ADD) Oriente, Occidente, Sur y Centro.', 1)
añadir_parrafo('En ' + mesesSTR[0].lower() + ', ' + mesesSTR[1].lower() + ' y ' + mesesSTR[2].lower() + ' de ' + str(Año) + ', se publicaron las liquidaciones ADD correspondientes a los meses de ' + MesesADD[2].lower() + ' de ' + str(AñosADD[8]) + ', ' + MesesADD[5].lower() + ' de ' + str(AñosADD[5]) + ' y ' + MesesADD[8].lower() + ' de ' + str(AñosADD[2]) + ' respectivamente.')
añadir_parrafo('En la gráfica 8, la gráfica 9, la gráfica 10 y la gráfica 11 se puede ver la evolución de los ingresos reconocidos e ingresos de la ADD oriente, occidente, sur y centro respectivamente, en cada uno de los niveles de tensión 1, 2 y 3.')

# Grafica ADD Oriente
doc.add_picture(r"Temp\ADD_Oriente.png")
añadir_parrafo('Fig 8, Ingreos ADD Oriente', 2)

# . Grafica ADD Occidente
doc.add_picture(r"Temp\ADD_Occidente.png")
añadir_parrafo('Fig 9, Ingreos ADD Occidente', 2)

# . Grafica ADD Centro
doc.add_picture(r"Temp\ADD_Centro.png")
añadir_parrafo('Fig 10, Ingreos ADD Centro', 2)

# . Grafica ADD Sur
doc.add_picture(r"Temp\ADD_Sur.png")
añadir_parrafo('Fig 11, Ingreos ADD Sur', 2)

## 5.SDL
añadir_parrafo('Liquidación SDL', 4)
añadir_parrafo('El Sistema de Distribución Local (SDL). Es el sistema de transporte de energía eléctrica compuesto por el conjunto de líneas y subestaciones, con sus equipos asociados, que operan a los Niveles de Tensión 3, 2 y 1 dedicados a la prestación del servicio en un mercado de comercialización.')
añadir_parrafo('Desde la grafica 12 hasta la grafica 16, se puede ver la evolución de los cargos asociados al nivel de tensión 1 para los 29 operadores de red que conforman el SDL durante el ' + Trimestre + ' del año ' + Año + '.')

# SDL Oriente
doc.add_picture(r"Temp\SDL_Cargos_Oriente.png")
añadir_parrafo('Fig 12, Evolución de cargo para nivel de tenión 1 para ORs asociados a ADD Oriente', 2)

# SDL Occidente
doc.add_picture(r"Temp\SDL_Cargos_Occidente.png")
añadir_parrafo('Fig 13, Evolución de cargo para nivel de tenión 1 para ORs asociados a ADD Occidente', 2)

# SDL Centro
doc.add_picture(r"Temp\SDL_Cargos_Centro.png")
añadir_parrafo('Fig 14, Evolución de cargo para nivel de tenión 1 para ORs asociados a ADD Centro', 2)

# SDL Sur
doc.add_picture(r"Temp\SDL_Cargos_Sur.png")
añadir_parrafo('Fig 15, Evolución de cargo para nivel de tenión 1 para ORs asociados a ADD Sur', 2)

# SDL sin ADDs
doc.add_picture(r"Temp\SDL_Cargos_sinADD.png")
añadir_parrafo('Fig 16, Evolución de cargo para nivel de tenión 1 para ORs sin ADD asociada', 2)

## 6.Calidad SDL
añadir_parrafo('Calidad del SDL', 4)
añadir_parrafo('Para el proceso de Calidad en el SDL se calculan los índices basados en la cantidad de incidentes presentados en el sistema dentro de un mercado de comercializacion especifico y la duracion de los mismos')

# . Grafica Calidad Inidcadores Globales
doc.add_picture(r"Temp\Calidad_SDL.png")
añadir_parrafo('Fig 17, Evolución de SAIDI y SAIFI a nivel pais.', 2)

## 7.CPROG
añadir_parrafo('Cargo y Liquidación CPROG', 4)
añadir_parrafo('Cargo por concepto del plan de reducción de pérdidas no técnicas que se traslada a los usuarios regulados y no regulados del mercado de comercialización.')

# . Grafica CPROG ADD Oriente
doc.add_picture(r"Temp\CPROG_Oriente.png")
añadir_parrafo('Fig 18, Evolución de cargo CPROG para ORs asociados a la ADD Oriente', 2)

# . Grafica CPROG ADD Occidente
doc.add_picture(r"Temp\CPROG_Occidente.png")
añadir_parrafo('Fig 19, Evolución de cargo CPROG para ORs asociados a la ADD Occidente', 2)

# . Grafica CPROG ADD Centro
doc.add_picture(r"Temp\CPROG_Centro.png")
añadir_parrafo('Fig 20, Evolución de cargo CPROG para ORs asociados a la ADD Centro', 2)

# . Grafica CPROG ADD Sur
doc.add_picture(r"Temp\CPROG_Sur.png")
añadir_parrafo('Fig 21, Evolución de cargo CPROG para ORs asociados a la ADD Sur', 2)

# . Grafica CPROG sin ADD
doc.add_picture(r"Temp\CPROG_sinADD.png")
añadir_parrafo('Fig 22, Evolución de cargo CPROG para ORs sin ADD asociada', 2)

## Document Export

doc.save(r"{name}.docx".format(name='Informe Mandatario ' + str(Trimestre) + ' año '+ str(Año)))
print('Done Exports!')


